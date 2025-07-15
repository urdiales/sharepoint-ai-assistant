"""
File processing utilities for the SharePoint AI Assistant.
Provides enhanced file preview and processing capabilities with error handling.
"""

import io
from typing import Optional, Dict, Any
from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document
import openpyxl

from ..core import (
    SharePointConstants,
    FileOperationError,
    FileTooLargeError,
    InvalidFileTypeError,
    get_logger,
    log_performance
)
from .validation import validate_filename, validate_file_extension, validate_file_size

# Get logger for this module
logger = get_logger("file_utils")


def preview_pdf(file_bytes: io.BytesIO, max_pages: int = 1) -> str:
    """
    Extract text from PDF file with enhanced error handling.
    
    Args:
        file_bytes: BytesIO object containing PDF data
        max_pages: Maximum number of pages to process
    
    Returns:
        Extracted text from PDF
    
    Raises:
        FileOperationError: If PDF processing fails
    """
    if not file_bytes:
        raise FileOperationError("PDF file data is empty")
    
    try:
        with log_performance(logger, "PDF text extraction"):
            # Check file size
            file_bytes.seek(0, 2)  # Seek to end
            file_size = file_bytes.tell()
            file_bytes.seek(0)  # Reset to beginning
            
            validate_file_size(file_size, SharePointConstants.MAX_PREVIEW_SIZE)
            
            # Extract text from PDF
            with pdfplumber.open(file_bytes) as pdf:
                if not pdf.pages:
                    logger.warning("PDF file contains no pages")
                    return "PDF file appears to be empty"
                
                text_parts = []
                pages_to_process = min(max_pages, len(pdf.pages))
                
                logger.info(f"Processing {pages_to_process} pages from PDF")
                
                for i in range(pages_to_process):
                    try:
                        page = pdf.pages[i]
                        page_text = page.extract_text()
                        
                        if page_text:
                            text_parts.append(page_text)
                        else:
                            logger.warning(f"No text found on page {i + 1}")
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {i + 1}: {e}")
                        continue
                
                if not text_parts:
                    return "No readable text found in PDF"
                
                # Join all text and limit length
                full_text = "\n\n".join(text_parts)
                
                if len(full_text) > SharePointConstants.PREVIEW_TEXT_LIMIT:
                    full_text = full_text[:SharePointConstants.PREVIEW_TEXT_LIMIT] + "..."
                    logger.info(f"PDF text truncated to {SharePointConstants.PREVIEW_TEXT_LIMIT} characters")
                
                logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
                return full_text
                
    except Exception as e:
        logger.error(f"Failed to process PDF file: {e}")
        raise FileOperationError(f"PDF processing failed: {str(e)}")


def preview_docx(file_bytes: io.BytesIO) -> str:
    """
    Extract text from DOCX file with enhanced error handling.
    
    Args:
        file_bytes: BytesIO object containing DOCX data
    
    Returns:
        Extracted text from DOCX
    
    Raises:
        FileOperationError: If DOCX processing fails
    """
    if not file_bytes:
        raise FileOperationError("DOCX file data is empty")
    
    try:
        with log_performance(logger, "DOCX text extraction"):
            # Check file size
            file_bytes.seek(0, 2)  # Seek to end
            file_size = file_bytes.tell()
            file_bytes.seek(0)  # Reset to beginning
            
            validate_file_size(file_size, SharePointConstants.MAX_PREVIEW_SIZE)
            
            # Extract text from DOCX
            doc = Document(file_bytes)
            
            if not doc.paragraphs:
                logger.warning("DOCX file contains no paragraphs")
                return "DOCX file appears to be empty"
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                return "No readable text found in DOCX"
            
            # Join all text and limit length
            full_text = "\n".join(text_parts)
            
            if len(full_text) > SharePointConstants.PREVIEW_TEXT_LIMIT:
                full_text = full_text[:SharePointConstants.PREVIEW_TEXT_LIMIT] + "..."
                logger.info(f"DOCX text truncated to {SharePointConstants.PREVIEW_TEXT_LIMIT} characters")
            
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
            
    except Exception as e:
        logger.error(f"Failed to process DOCX file: {e}")
        raise FileOperationError(f"DOCX processing failed: {str(e)}")


def preview_xlsx(file_bytes: io.BytesIO, max_rows: int = 10) -> pd.DataFrame:
    """
    Extract data from XLSX file with enhanced error handling.
    
    Args:
        file_bytes: BytesIO object containing XLSX data
        max_rows: Maximum number of rows to return
    
    Returns:
        DataFrame containing Excel data
    
    Raises:
        FileOperationError: If XLSX processing fails
    """
    if not file_bytes:
        raise FileOperationError("XLSX file data is empty")
    
    try:
        with log_performance(logger, "XLSX data extraction"):
            # Check file size
            file_bytes.seek(0, 2)  # Seek to end
            file_size = file_bytes.tell()
            file_bytes.seek(0)  # Reset to beginning
            
            validate_file_size(file_size, SharePointConstants.MAX_PREVIEW_SIZE)
            
            # Load workbook
            wb = openpyxl.load_workbook(file_bytes, read_only=True, data_only=True)
            
            if not wb.worksheets:
                logger.warning("XLSX file contains no worksheets")
                return pd.DataFrame({"Error": ["No worksheets found in file"]})
            
            # Get the active worksheet
            ws = wb.active
            
            if not ws:
                logger.warning("XLSX file has no active worksheet")
                return pd.DataFrame({"Error": ["No active worksheet found"]})
            
            # Extract data with row limit
            data = []
            headers = None
            rows_processed = 0
            
            for row in ws.iter_rows(values_only=True):
                if rows_processed >= max_rows + 1:  # +1 for header
                    break
                
                # Skip completely empty rows
                if all(cell is None or str(cell).strip() == "" for cell in row):
                    continue
                
                # Convert None values to empty strings and handle data types
                processed_row = []
                for cell in row:
                    if cell is None:
                        processed_row.append("")
                    else:
                        processed_row.append(str(cell))
                
                if headers is None:
                    headers = processed_row
                    logger.debug(f"Found headers: {headers}")
                else:
                    data.append(processed_row)
                    rows_processed += 1
            
            if not headers:
                logger.warning("No data found in XLSX file")
                return pd.DataFrame({"Error": ["No data found in file"]})
            
            if not data:
                logger.warning("No data rows found in XLSX file")
                return pd.DataFrame(columns=headers)
            
            # Create DataFrame
            df = pd.DataFrame(data, columns=headers)
            
            # Clean up column names
            df.columns = [str(col).strip() if col else f"Column_{i}" for i, col in enumerate(df.columns)]
            
            logger.info(f"Successfully extracted {len(df)} rows and {len(df.columns)} columns from XLSX")
            return df
            
    except Exception as e:
        logger.error(f"Failed to process XLSX file: {e}")
        raise FileOperationError(f"XLSX processing failed: {str(e)}")


def get_file_info(file_bytes: io.BytesIO, filename: str) -> Dict[str, Any]:
    """
    Get comprehensive information about a file.
    
    Args:
        file_bytes: BytesIO object containing file data
        filename: Name of the file
    
    Returns:
        Dictionary with file information
    
    Raises:
        FileOperationError: If file analysis fails
    """
    try:
        # Validate filename
        validated_filename = validate_filename(filename)
        file_extension = validate_file_extension(validated_filename)
        
        # Get file size
        file_bytes.seek(0, 2)  # Seek to end
        file_size = file_bytes.tell()
        file_bytes.seek(0)  # Reset to beginning
        
        # Basic file info
        file_info = {
            "filename": validated_filename,
            "extension": file_extension,
            "size_bytes": file_size,
            "size_mb": round(file_size / (1024 * 1024), 2),
            "is_supported": file_extension in SharePointConstants.SUPPORTED_FILE_TYPES,
            "can_preview": file_size <= SharePointConstants.MAX_PREVIEW_SIZE
        }
        
        # Add type-specific information
        if file_extension == ".pdf":
            try:
                with pdfplumber.open(file_bytes) as pdf:
                    file_info["page_count"] = len(pdf.pages)
                    file_info["has_text"] = any(page.extract_text() for page in pdf.pages[:3])
            except Exception as e:
                logger.warning(f"Could not analyze PDF structure: {e}")
                file_info["analysis_error"] = str(e)
        
        elif file_extension == ".docx":
            try:
                doc = Document(file_bytes)
                file_info["paragraph_count"] = len(doc.paragraphs)
                file_info["table_count"] = len(doc.tables)
                file_info["has_content"] = bool(doc.paragraphs or doc.tables)
            except Exception as e:
                logger.warning(f"Could not analyze DOCX structure: {e}")
                file_info["analysis_error"] = str(e)
        
        elif file_extension == ".xlsx":
            try:
                wb = openpyxl.load_workbook(file_bytes, read_only=True)
                file_info["worksheet_count"] = len(wb.worksheets)
                if wb.active:
                    file_info["has_data"] = wb.active.max_row > 1
                    file_info["max_row"] = wb.active.max_row
                    file_info["max_column"] = wb.active.max_column
            except Exception as e:
                logger.warning(f"Could not analyze XLSX structure: {e}")
                file_info["analysis_error"] = str(e)
        
        # Reset file position
        file_bytes.seek(0)
        
        logger.debug(f"File analysis complete: {file_info}")
        return file_info
        
    except Exception as e:
        logger.error(f"Failed to analyze file {filename}: {e}")
        raise FileOperationError(f"File analysis failed: {str(e)}")


def format_file_size(size_bytes: int) -> str:
    """
    Format file size in human-readable format.
    
    Args:
        size_bytes: File size in bytes
    
    Returns:
        Formatted file size string
    """
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    elif size_bytes < 1024 * 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.1f} MB"
    else:
        return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"


def is_file_type_supported(filename: str) -> bool:
    """
    Check if file type is supported for preview.
    
    Args:
        filename: Name of the file
    
    Returns:
        True if file type is supported
    """
    try:
        extension = validate_file_extension(filename, SharePointConstants.SUPPORTED_FILE_TYPES)
        return True
    except Exception:
        return False


def get_preview_function(filename: str):
    """
    Get the appropriate preview function for a file type.
    
    Args:
        filename: Name of the file
    
    Returns:
        Preview function for the file type
    
    Raises:
        InvalidFileTypeError: If file type is not supported
    """
    try:
        extension = validate_file_extension(filename, SharePointConstants.SUPPORTED_FILE_TYPES)
        
        preview_functions = {
            ".pdf": preview_pdf,
            ".docx": preview_docx,
            ".xlsx": preview_xlsx
        }
        
        if extension not in preview_functions:
            raise InvalidFileTypeError(f"No preview function available for {extension}")
        
        return preview_functions[extension]
        
    except Exception as e:
        raise InvalidFileTypeError(f"Cannot determine preview function: {str(e)}")
